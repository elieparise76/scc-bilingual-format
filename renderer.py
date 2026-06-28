"""Phase 4 — Renderer.

Génère le document Word (.docx) bilingue à deux colonnes à partir des
sections alignées.

Structure produite :
- en-tête courant **alterné par page** (anglais sur les pages impaires, français
  sur les paires) : « <nom de la décision en italique> — <référence> — <juge
  rédacteur et son rôle sur cette page> », suivi d'une ligne horizontale ;
- pied de page : numéro de page ;
- page de garde (titre + référence) et table des matières des sections ;
- corps : par opinion, un tableau deux colonnes (langue de gauche / langue de
  droite), une ligne par paire de paragraphes, alignées en haut de cellule ;
- bandeaux (sans couleur) au début de chaque opinion ;
- lignes de sous-titres (« II. Contexte », « A. … ») au-dessus des paragraphes.

L'en-tête « juge de cette page » s'appuie sur le champ Word **STYLEREF** : à
chaque début d'opinion on insère deux paragraphes marqueurs *masqués* (un par
langue, styles `OpinionRefEN`/`OpinionRefFR`) contenant « Auteur (rôle) » ;
STYLEREF affiche dans l'en-tête le premier marqueur présent sur la page, ou à
défaut le dernier avant la page — soit exactement l'opinion qui débute (ou se
poursuit) sur la page. Marqueurs au niveau du corps (hors tableau) pour que
STYLEREF les retrouve de façon fiable.

Réglages globaux : Times New Roman partout, interligne simple, texte justifié,
marges 0,4 po, bordures de tableau blanches, espace accru entre les colonnes.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from itertools import zip_longest
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

from citation_link import emit_runs as _emit_citation_runs
from models import AlignedSection, Paragraph, SectionType, TextBlock

_FONT = "Times New Roman"
_PAGE_W = Inches(8.5)
_MARGIN = Inches(0.4)
_CONTENT_W = Emu(_PAGE_W - _MARGIN * 2)  # largeur utile entre marges (7,7 po)
_COL_W = Emu(_CONTENT_W // 2)            # colonnes strictement égales (3,85 po)
_ZERO = Inches(0)
_CELL_GAP = Inches(0.1)  # demi-espace au centre (marge interne du côté centre)
_CELL_TB = Inches(0.03)
# Le nom de cause est tronqué pour tenir sur une ligne (sinon il déborde les
# taquets et la citation ne peut pas atteindre la droite).
_HEADER_TITLE_MAX = 56
# Page 1 (couverture) : pas de juge au centre → beaucoup plus de place pour le
# nom de cause avant de devoir le tronquer.
_HEADER_TITLE_MAX_FIRST = 90

# Texte « en retrait » (citations en bloc, listes) : corps légèrement plus petit
# que le corps courant (11 pt) et retrait gauche croissant selon le niveau.
_INDENT_FONT = Pt(10)
_INDENT_BASE = Inches(0.05)   # retrait du 1er niveau (au-delà du bord de cellule)
_INDENT_STEP = Inches(0.25)   # retrait additionnel par niveau

# Libellés de rôle de section (fr, en) pour bandeaux et table des matières.
_ROLE_LABELS = {
    SectionType.MAJORITY: ("Motifs de la majorité", "Majority reasons"),
    SectionType.CONCURRING: ("Motifs concordants", "Concurring reasons"),
    SectionType.DISSENT: ("Motifs dissidents", "Dissenting reasons"),
    SectionType.HEADNOTES: ("Sommaire", "Headnotes"),
    SectionType.OTHER: ("Autres motifs", "Other reasons"),
}
_COURT_AUTHORS = {"La Cour", "The Court"}

# Libellés de la page de garde (fr, en).
_HEARD_LABEL = {"fr": "Appel entendu", "en": "Appeal heard"}
_DATE_LABEL = {"fr": "Jugement rendu", "en": "Judgment rendered"}
_TOC_LABEL = {"fr": "Table des matières", "en": "Table of Contents"}
_DOCKET_LABEL = {"fr": "Dossier no", "en": "Case No."}
_CORAM_LABEL = {"fr": "CORAM :", "en": "CORAM:"}
# Avis de non-officialité (bas de la page de garde) : le document est dérivé des
# PDF officiels mais remis en forme (côte à côte, réaligné) → sans valeur officielle.
_NOTICE = {
    "fr": "Version non officielle, remise en forme à partir des motifs officiels "
          "de la Cour suprême du Canada. Consultez la version officielle sur "
          "decisions.scc-csc.ca.",
    "en": "Unofficial version, reformatted from the official reasons of the Supreme "
          "Court of Canada. Consult the official version at decisions.scc-csc.ca.",
}
# Taquet droit (avec points de conduite) pour la table des matières, dans une
# cellule = largeur de texte de la cellule (_COL_W moins le gap intérieur).
_TOC_TAB = Emu(_COL_W - _CELL_GAP)


@dataclass
class DocMetadata:
    """Métadonnées de page de garde / en-tête, dans les deux langues."""

    title_fr: str
    title_en: str
    citation_fr: str
    citation_en: str
    hearing_fr: str = ""
    hearing_en: str = ""
    date_fr: str = ""
    date_en: str = ""
    appeal_fr: str = ""
    appeal_en: str = ""
    catchwords_fr: str = ""
    catchwords_en: str = ""
    held_fr: str = ""
    held_en: str = ""
    parties_fr: List[Tuple[str, str]] = field(default_factory=list)
    parties_en: List[Tuple[str, str]] = field(default_factory=list)
    coram_fr: List[str] = field(default_factory=list)
    coram_en: List[str] = field(default_factory=list)
    docket_fr: str = ""
    docket_en: str = ""
    lang_order: str = "en"  # "en" → EN à gauche (défaut) ; "fr" → FR à gauche

    def title(self, lang: str) -> str:
        return self.title_fr if lang == "fr" else self.title_en

    def citation(self, lang: str) -> str:
        return self.citation_fr if lang == "fr" else self.citation_en

    def hearing(self, lang: str) -> str:
        return self.hearing_fr if lang == "fr" else self.hearing_en

    def date(self, lang: str) -> str:
        return self.date_fr if lang == "fr" else self.date_en

    def appeal(self, lang: str) -> str:
        return self.appeal_fr if lang == "fr" else self.appeal_en

    def catchwords(self, lang: str) -> str:
        return self.catchwords_fr if lang == "fr" else self.catchwords_en

    def held(self, lang: str) -> str:
        return self.held_fr if lang == "fr" else self.held_en

    def docket(self, lang: str) -> str:
        return self.docket_fr if lang == "fr" else self.docket_en

    def parties(self, lang: str) -> List[Tuple[str, str]]:
        return self.parties_fr if lang == "fr" else self.parties_en

    def coram(self, lang: str) -> List[str]:
        return self.coram_fr if lang == "fr" else self.coram_en


# --------------------------------------------------------------------------- #
# Helpers bas niveau (OOXML)
# --------------------------------------------------------------------------- #
def _add_page_field(paragraph) -> None:
    """Insère un champ « PAGE » (numéro de page dynamique)."""
    _add_field(paragraph, "PAGE", cached="1")


def _add_field(paragraph, instruction: str, cached: Optional[str] = None):
    """Insère un champ Word ; `cached` = résultat affiché avant rafraîchissement.

    Sans résultat mis en cache, Word affiche le champ **vide** tant qu'il n'est
    pas recalculé (d'où l'impression que le juge « manque » dans l'en-tête).
    """
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(begin)
    run._r.append(instr)
    if cached is None:
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(end)
        return run
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    result = paragraph.add_run(cached)  # valeur en cache (sera mise à jour)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    return result


def _enable_odd_even(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def _enable_update_fields(doc: Document) -> None:
    """Demande à Word de recalculer tous les champs à l'ouverture."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def _para_bottom_border(paragraph) -> None:
    """Ligne horizontale sous le paragraphe (utilisé pour l'en-tête)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _set_base_font(doc: Document) -> None:
    """Times New Roman + interligne simple sur les styles utilisés."""
    for name in ("Normal", "Header", "Footer", "List Bullet"):
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = _FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), _FONT)
        # Les styles Header/Footer ont des taquets par défaut (centre 3,25 po,
        # droite 6,5 po) qui, fusionnés avec les nôtres, captent la citation
        # avant le bord droit. On les retire pour que seuls nos taquets jouent.
        if name in ("Header", "Footer"):
            pPr = style.element.find(qn("w:pPr"))
            tabs = pPr.find(qn("w:tabs")) if pPr is not None else None
            if tabs is not None:
                pPr.remove(tabs)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _set_table_borders_white(table) -> None:
    """Rend toutes les bordures du tableau blanches (lignes invisibles)."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "FFFFFF")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_margins(cell, left, right) -> None:
    """Marges internes d'une cellule. Pour que le texte aille bord à bord (et
    s'aligne sur le trait de l'en-tête), la marge **extérieure est nulle** ; seul
    le côté **centre** porte `_CELL_GAP`, ce qui crée l'espace inter-colonnes."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", _CELL_TB), ("left", left), ("bottom", _CELL_TB), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(int(val.twips)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _col_margins(col_index: int) -> tuple:
    """(marge gauche, marge droite) interne : 0 à l'extérieur, gap au centre."""
    return (_ZERO, _CELL_GAP) if col_index == 0 else (_CELL_GAP, _ZERO)


def _create_ref_styles(doc: Document) -> None:
    """Styles « invisibles » portant « Auteur (rôle) » pour les champs STYLEREF.

    Rendus invisibles par **blanc + 1 pt** (et non `w:vanish`) : le texte masqué
    n'est pas mis en page, or STYLEREF dépend de la pagination — un marqueur
    masqué est introuvable et le champ reste vide. Un marqueur blanc minuscule
    occupe une position de page tout en restant invisible.
    """
    for name in ("OpinionRefEN", "OpinionRefFR"):
        if name in doc.styles:
            continue
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles["Normal"]
        st.font.size = Pt(1)
        st.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        pf = st.paragraph_format
        pf.line_spacing = Pt(1)  # hauteur de ligne minimale
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


# --------------------------------------------------------------------------- #
# Libellés (rôles, auteurs)
# --------------------------------------------------------------------------- #
def _lead_author(author: str) -> str:
    """« La juge Martin (avec l'accord…) » → « La juge Martin »."""
    return author.split("(")[0].strip() if author else ""


def _section_label(section: AlignedSection, lang: str, unanimous: bool) -> str:
    """Libellé de bandeau/TOC : « Motifs de la majorité — La juge Côté »."""
    if unanimous:
        role = "Motifs de la Cour" if lang == "fr" else "Reasons of the Court"
    else:
        fr, en = _ROLE_LABELS.get(section.type, _ROLE_LABELS[SectionType.OTHER])
        role = fr if lang == "fr" else en
    author = _lead_author(section.author_fr if lang == "fr" else section.author_en)
    if author in _COURT_AUTHORS:  # éviter « Motifs de la Cour — La Cour »
        return role
    return f"{role} — {author}" if author else role


# --------------------------------------------------------------------------- #
# En-tête / pied de page
# --------------------------------------------------------------------------- #
def _header_case(title: str, max_len: int = _HEADER_TITLE_MAX) -> str:
    """Tronque le nom de cause (au mot) pour qu'il tienne sur une ligne."""
    if len(title) <= max_len:
        return title
    return title[:max_len].rsplit(" ", 1)[0] + "…"


def _fill_header(
    paragraph, title: str, citation: str, ref_style: str, cached: str,
    *, with_judge: bool = True,
) -> None:
    """En-tête à trois zones par taquets : n° de page (gauche), *cause* — juge
    (centre, cause en italique), citation (droite). Le juge est un champ
    STYLEREF (varie selon la page) ; sans rôle (majoritaire/unanime…).

    `with_judge=False` (page 1 / couverture) : le centre n'affiche que le nom de
    cause, sans « — » ni champ STYLEREF — la couverture précède toute opinion,
    donc aucun juge rédacteur, et la page 1 n'utilise plus STYLEREF.

    Aucun retrait de paragraphe : les taquets se mesurent depuis la marge et le
    taquet droit à `_CONTENT_W` place la citation **au bord droit**. Le texte des
    colonnes va aussi bord à bord (marges extérieures nulles), donc le trait du
    bas s'aligne dessus.
    """
    for run in list(paragraph.runs):  # vider sans laisser de run vide
        run._r.getparent().remove(run._r)
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.tab_stops.add_tab_stop(_COL_W, WD_TAB_ALIGNMENT.CENTER)      # centre
    pf.tab_stops.add_tab_stop(_CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)   # bord droit

    # Page 1 (sans juge) : titre tronqué bien plus tard (plus de place au centre).
    title_max = _HEADER_TITLE_MAX if with_judge else _HEADER_TITLE_MAX_FIRST
    _add_page_field(paragraph)                          # gauche : n° de page
    paragraph.add_run("\t")
    paragraph.add_run(_header_case(title, title_max)).italic = True  # centre : cause — juge
    if with_judge:
        paragraph.add_run(" — ")
        _add_field(paragraph, f'STYLEREF "{ref_style}"', cached=cached)
    paragraph.add_run("\t")
    paragraph.add_run(citation)                         # droite : citation
    for run in paragraph.runs:
        run.font.size = Pt(9)
    _para_bottom_border(paragraph)


def _setup_headers_footers(
    doc: Document, meta: DocMetadata, first: AlignedSection
) -> None:
    sec = doc.sections[0]
    # Page 1 (garde) avec un en-tête propre, SANS juge ; en-tête alterné (avec
    # juge) à partir de la page 2.
    sec.different_first_page_header_footer = True
    first_header = sec.first_page_header
    odd_header = sec.header
    even_header = sec.even_page_header
    first_header.is_linked_to_previous = False
    odd_header.is_linked_to_previous = False
    even_header.is_linked_to_previous = False
    # Page 1 (impaire) → anglais, sans juge (couverture, pas d'opinion → pas de
    # STYLEREF, donc aucun risque d'« Error! No text of specified style »).
    _fill_header(
        first_header.paragraphs[0], meta.title_en, meta.citation_en,
        "OpinionRefEN", _lead_author(first.author_en), with_judge=False,
    )
    # Pages impaires (sauf 1) → anglais ; paires → français. Après la garde
    # (page 1), page 2 est paire → français, page 3 impaire → anglais.
    _fill_header(
        odd_header.paragraphs[0], meta.title_en, meta.citation_en,
        "OpinionRefEN", _lead_author(first.author_en),
    )
    _fill_header(
        even_header.paragraphs[0], meta.title_fr, meta.citation_fr,
        "OpinionRefFR", _lead_author(first.author_fr),
    )
    # Pas de numéro de page en pied : il est désormais dans l'en-tête.


# --------------------------------------------------------------------------- #
# Page de garde bilingue (identité, mots-clés, table des matières)
# --------------------------------------------------------------------------- #
def _heading_level(heading: str) -> int:
    """Niveau de plan d'un sous-titre, pour l'indentation de la table des matières."""
    h = heading.lstrip()
    if re.match(r"[IVXL]+\.", h):
        return 0
    if re.match(r"[A-Z]\.", h):
        return 1
    if re.match(r"\(\d+\)", h):
        return 2
    return 3  # (a), (i)…


def _section_toc(section: AlignedSection, lang: str) -> List[tuple]:
    """[(sous-titre, n° de paragraphe, niveau)] d'une opinion, pour une langue."""
    out: List[tuple] = []
    for pair in section.pairs:
        para = pair.fr if lang == "fr" else pair.en
        if para:
            out.extend((h, pair.number, _heading_level(h)) for h in para.headings)
    return out


def _new_aligned_table(doc: Document):
    """Tableau 2 colonnes égales, bordures blanches (page de garde / opinions)."""
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders_white(table)
    return table


def _set_run(run, size: int, bold: bool = False, italic: bool = False):
    run.font.size = Pt(size)
    run.bold, run.italic = bold, italic
    return run


def _bi_row(
    table, lang_left, lang_right, text_fn, *, size=9, bold=False, italic=False,
    align=WD_ALIGN_PARAGRAPH.LEFT, before=0,
) -> None:
    """Ajoute une ligne (gauche/droite) ; chaque côté est `text_fn(lang)`.

    Une ligne par élément → les deux langues démarrent au même niveau (comme
    les paragraphes du corps). Ligne omise si les deux côtés sont vides."""
    left, right = text_fn(lang_left), text_fn(lang_right)
    if not left and not right:
        return
    row = table.add_row()
    for i, (cell, text) in enumerate(((row.cells[0], left), (row.cells[1], right))):
        _set_cell_margins(cell, *_col_margins(i))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = align
        if before:
            p.paragraph_format.space_before = Pt(before)
        _set_run(p.add_run(text), size, bold=bold, italic=italic)


def _add_party_row(table, sides: List[tuple], *, before: int = 0) -> None:
    """Ligne de partie bilingue : « **Rôle** — noms » (rôle en gras).

    `sides` = [(rôle_gauche, noms_gauche), (rôle_droite, noms_droite)]. Si le
    rôle est absent, n'affiche que les noms (sans tiret). Ligne omise si les
    deux côtés sont vides."""
    if all(not role and not names for role, names in sides):
        return
    row = table.add_row()
    for i, (cell, (role, names)) in enumerate(zip((row.cells[0], row.cells[1]), sides)):
        _set_cell_margins(cell, *_col_margins(i))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if before:
            p.paragraph_format.space_before = Pt(before)
        if role:
            _set_run(p.add_run(role), 9, bold=True)
            if names:
                _set_run(p.add_run(f" — {names}"), 9)
        elif names:
            _set_run(p.add_run(names), 9)


def _add_toc_row(table, left: tuple, right: tuple) -> None:
    """Ligne de TDM ; left/right = (libellé, texte_droite, niveau).

    Le libellé est indenté selon le niveau ; le texte de droite (n° ou plage)
    est aligné à droite avec des points de conduite."""
    row = table.add_row()
    for i, (cell, item) in enumerate(((row.cells[0], left), (row.cells[1], right))):
        _set_cell_margins(cell, *_col_margins(i))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        label, right_text, level = item
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.left_indent = Emu(Inches(0.13) * level)
        _set_run(p.add_run(label), 9)
        if right_text:
            pf.tab_stops.add_tab_stop(_TOC_TAB, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.add_run("\t")
            _set_run(p.add_run(str(right_text)), 9)


def _add_contents_toc(table, lang_left, lang_right, entries_fn) -> None:
    """Table des matières : ligne de titre puis lignes appariées gauche/droite."""
    _bi_row(table, lang_left, lang_right, lambda lang: _TOC_LABEL[lang],
            size=10, bold=True, before=10)
    left_items = entries_fn(lang_left)
    right_items = entries_fn(lang_right)
    for left, right in zip_longest(left_items, right_items, fillvalue=("", "", 0)):
        _add_toc_row(table, left, right)


def _add_front_matter(
    doc: Document, meta: DocMetadata, sections: List[AlignedSection], lang_left: str
) -> None:
    """Page de garde bilingue : un élément par ligne, deux colonnes alignées.

    Ordre, de haut en bas : **avis de non-officialité** (en tête), **nom de
    cause** (grand titre, comme avant — il figure aussi dans l'en-tête courant),
    citation, n° de greffe, date d'audition, date du jugement, **parties**,
    **CORAM**, puis mention d'appel, mots-clés/tags, mention « Held / Arrêt » et
    la table des matières des **motifs** (rôle + auteur + plage de paragraphes)."""
    lang_right = "en" if lang_left == "fr" else "fr"
    unanimous = len(sections) == 1
    table = _new_aligned_table(doc)
    center = WD_ALIGN_PARAGRAPH.CENTER

    def dated(label, getter):
        return lambda lang: (
            f"{label[lang]}{' : ' if lang == 'fr' else ': '}{getter(lang)}"
            if getter(lang) else ""
        )

    # Avis de non-officialité en TÊTE de la couverture.
    _bi_row(table, lang_left, lang_right, lambda lang: _NOTICE[lang],
            italic=True, align=center)
    # Titre : pas d'espace avant (suite directe de l'avis).
    _bi_row(table, lang_left, lang_right, meta.title, size=12, bold=True, align=center)
    _bi_row(table, lang_left, lang_right, meta.citation, size=11, bold=True, align=center)
    _bi_row(table, lang_left, lang_right,
            lambda lang: f"{_DOCKET_LABEL[lang]} {meta.docket(lang)}" if meta.docket(lang) else "",
            align=center)
    _bi_row(table, lang_left, lang_right, dated(_HEARD_LABEL, meta.hearing),
            align=center, before=6)
    _bi_row(table, lang_left, lang_right, dated(_DATE_LABEL, meta.date), align=center)

    # CORAM (avant les parties)
    _bi_row(table, lang_left, lang_right,
            lambda lang: (f"{_CORAM_LABEL[lang]} {', '.join(meta.coram(lang))}"
                          if meta.coram(lang) else ""),
            before=6)

    # Mention d'appel : avant les parties.
    _bi_row(table, lang_left, lang_right, meta.appeal, italic=True, align=center, before=6)

    # Parties : « **Rôle** — noms » (rôle en gras), alignées FR/EN par position.
    parties_l = meta.parties(lang_left)
    parties_r = meta.parties(lang_right)
    for i, (pl, pr) in enumerate(zip_longest(parties_l, parties_r, fillvalue=("", ""))):
        nom_l, role_l = pl
        nom_r, role_r = pr
        _add_party_row(table, [(role_l, nom_l), (role_r, nom_r)],
                       before=(8 if i == 0 else 2))
    _bi_row(table, lang_left, lang_right, meta.catchwords,
            italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=10)

    # HELD / ARRÊT : premier mot en majuscules, toute la ligne en gras.
    def held_upper(lang):
        text = meta.held(lang)
        return re.sub(r'^(\w+)', lambda m: m.group(1).upper(), text) if text else ""

    _bi_row(table, lang_left, lang_right, held_upper,
            bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=8)

    # Table des matières = les motifs (rôle + auteur + plage), même si unanime.
    def reasons(lang):
        items = []
        for sec in sections:
            nums = [p.number for p in sec.pairs]
            rng = f"[{nums[0]}]-[{nums[-1]}]" if nums else ""
            items.append((_section_label(sec, lang, unanimous), rng, 0))
        return items

    _add_contents_toc(table, lang_left, lang_right, reasons)


# --------------------------------------------------------------------------- #
# Corps : marqueurs, bandeaux, tableaux
# --------------------------------------------------------------------------- #
def _add_ref_markers(doc: Document, section: AlignedSection) -> None:
    """Marqueurs (invisibles) pour STYLEREF en début d'opinion : juge seul."""
    doc.add_paragraph(_lead_author(section.author_en), style="OpinionRefEN")
    doc.add_paragraph(_lead_author(section.author_fr), style="OpinionRefFR")


def _add_banner_row(
    table, section: AlignedSection, lang_left: str, unanimous: bool
) -> None:
    """Bandeau de début d'opinion : mention d'attribution verbatim, **par côté**.

    Ligne à deux colonnes (alignées sur le corps) : à gauche la mention de la
    langue de gauche (« Le jugement suivant a été rendu par / LA COUR — »), à
    droite celle de la langue de droite. Repli sur un libellé synthétique si
    la mention manque.
    """
    lang_right = "en" if lang_left == "fr" else "fr"
    left = section.lead_in_fr if lang_left == "fr" else section.lead_in_en
    right = section.lead_in_en if lang_left == "fr" else section.lead_in_fr
    left = left or _section_label(section, lang_left, unanimous)
    right = right or _section_label(section, lang_right, unanimous)

    row = table.add_row()
    for i, (cell, text) in enumerate(((row.cells[0], left), (row.cells[1], right))):
        cell.width = _COL_W
        _set_cell_margins(cell, *_col_margins(i))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(text).bold = True  # « \n » → saut de ligne (préambule / auteur)


def _emit_runs(para, runs, size=None) -> None:
    """Écrit les fragments stylés (italique/gras) dans un paragraphe ; `size`
    force le corps (texte en retrait, plus petit).

    Seul point d'écriture des runs du corps (prose et blocs en retrait) : on y
    transforme chaque référence neutre CSC (« AAAA SCC/CSC N ») en hyperlien
    CanLII, sans toucher au reste du texte ni à sa mise en forme. Voir
    `citation_link`."""
    _emit_citation_runs(para, runs, size)


def _style_indent(para, indent: int) -> None:
    """Met en forme un paragraphe en retrait (citation/liste) : retrait gauche
    selon le niveau, léger espacement vertical. Justifié comme le corps."""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent >= 1:
        pf = para.paragraph_format
        pf.left_indent = Emu(_INDENT_BASE + _INDENT_STEP * indent)
        pf.space_before = Pt(4)
        pf.space_after = Pt(2)


def _write_cell(cell, paragraph: Optional[Paragraph], number: int, col: int) -> None:
    cell.width = _COL_W
    _set_cell_margins(cell, *_col_margins(col))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if paragraph is None:
        run = p.add_run("—")
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return

    # Blocs = prose + retraits (citations/listes). Repli : un bloc plat (corps).
    blocks = paragraph.blocks or [TextBlock(runs=paragraph.runs, indent=0)]
    for i, block in enumerate(blocks):
        para = p if i == 0 else cell.add_paragraph()
        _style_indent(para, block.indent)
        if i == 0:
            para.add_run(f"[{number}] ").bold = True  # marqueur sur le 1er bloc
        size = _INDENT_FONT if block.indent >= 1 else None
        if block.runs:
            _emit_runs(para, block.runs, size)
        else:
            run = para.add_run(paragraph.text)
            if size is not None:
                run.font.size = size


def _add_heading_row(table, pair, lang_left: str) -> None:
    """Ligne de sous-titres au-dessus d'un paragraphe (gras, deux colonnes)."""
    left_para = pair.fr if lang_left == "fr" else pair.en
    right_para = pair.en if lang_left == "fr" else pair.fr
    left_h = left_para.headings if left_para else []
    right_h = right_para.headings if right_para else []
    if not left_h and not right_h:
        return
    row = table.add_row()
    for i, (cell, headings) in enumerate(((row.cells[0], left_h), (row.cells[1], right_h))):
        cell.width = _COL_W
        _set_cell_margins(cell, *_col_margins(i))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.add_run("\n".join(headings)).bold = True


def _add_opinion_table(
    doc: Document, section: AlignedSection, lang_left: str, unanimous: bool
) -> None:
    lang_right = "en" if lang_left == "fr" else "fr"
    table = _new_aligned_table(doc)
    _add_banner_row(table, section, lang_left, unanimous)

    # Table des matières propre à l'opinion — uniquement s'il y a des sous-titres.
    if _section_toc(section, lang_left) or _section_toc(section, lang_right):
        _add_contents_toc(
            table, lang_left, lang_right, lambda lang: _section_toc(section, lang)
        )

    for pair in section.pairs:
        _add_heading_row(table, pair, lang_left)
        row = table.add_row()
        left_para = pair.fr if lang_left == "fr" else pair.en
        right_para = pair.en if lang_left == "fr" else pair.fr
        _write_cell(row.cells[0], left_para, pair.number, 0)
        _write_cell(row.cells[1], right_para, pair.number, 1)


# --------------------------------------------------------------------------- #
# Point d'entrée
# --------------------------------------------------------------------------- #
def render_docx(
    sections: List[AlignedSection],
    metadata: DocMetadata,
    output_path: str,
) -> str:
    """Produit le fichier .docx bilingue.

    Args:
        sections: sections alignées (sortie de l'Aligner).
        metadata: titres/citations FR+EN, ordre des langues.
        output_path: chemin du fichier .docx à écrire.

    Returns:
        Le chemin du fichier écrit.
    """
    lang_left = metadata.lang_order
    unanimous = len(sections) == 1

    doc = Document()
    for s in doc.sections:
        s.page_width = _PAGE_W
        s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = _MARGIN
        s.left_margin = s.right_margin = _MARGIN

    _enable_odd_even(doc)
    _enable_update_fields(doc)
    _set_base_font(doc)
    _create_ref_styles(doc)
    _setup_headers_footers(doc, metadata, sections[0])

    # Marqueur précoce pour la 1re opinion : garantit que STYLEREF se résout dès
    # la page 1 (sinon « Error! No text of specified style »).
    _add_ref_markers(doc, sections[0])
    _add_front_matter(doc, metadata, sections, lang_left)
    doc.add_page_break()

    for section in sections:
        _add_ref_markers(doc, section)
        _add_opinion_table(doc, section, lang_left, unanimous)

    doc.save(output_path)
    return output_path
