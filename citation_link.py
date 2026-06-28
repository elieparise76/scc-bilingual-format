"""Liens hypertexte du corps : références neutres → CanLII, lois → Justice Canada.

Détecte dans le corps du document trois familles de citations et transforme chaque
occurrence en hyperlien cliquable :

1. **Références neutres** « AAAA SCC N » (anglais) / « AAAA CSC N » (français)
   → **CanLII**. URL **déterministe**, sans appel réseau : la règle de citation
   CanLII pose que « when a decision has a neutral citation assigned by the issuing
   court, the CanLII citation is entirely based upon the neutral citation ». Le
   docID = la citation, espaces retirés, en minuscules. D'où le patron :

     « 2024 SCC 5 » → https://www.canlii.org/en/ca/scc/doc/2024/2024scc5/2024scc5.html
     « 2024 CSC 5 » → https://www.canlii.org/fr/ca/csc/doc/2024/2024csc5/2024csc5.html

   ⚠️ CanLII renvoie 403 à tout accès automatisé (anti-bot) : l'URL ne peut donc
   pas être validée par requête HTTP. C'est attendu — elle est garantie par la
   règle de citation ci-dessus et reste valide au clic humain dans Word.

2. **Lois fédérales des Lois révisées du Canada (1985)** « R.S.C. 1985, c. C-50 »
   (anglais) / « L.R.C. 1985, c. C-50 » (français) → **Justice Canada**
   (laws-lois.justice.gc.ca), le site officiel. URL **déterministe** : pour une loi
   des L.R.C. 1985, le **chapitre cité est le slug d'URL** (R.S.C. → /eng/acts/,
   L.R.C. → /fra/lois/) :

     « R.S.C. 1985, c. C-50 » → https://laws-lois.justice.gc.ca/eng/acts/C-50/
     « L.R.C. 1985, c. C-50 » → https://laws-lois.justice.gc.ca/fra/lois/C-50/

   Contrairement à CanLII, Justice Canada ne bloque pas les bots : l'URL est donc
   **vérifiable par requête HTTP** (utilisé pour garantir zéro lien mort).
   **Limité aux R.S.C./L.R.C. 1985** : les lois annuelles « S.C. 2012, c. 1 » et
   les suppléments « c. 1 (5e suppl.) » n'ont **pas** de slug déterministe par
   chapitre (l'Income Tax Act « R.S.C. 1985, c. 1 (5th Supp.) » a le slug réel
   « I-3.3 », pas « 1 » ; « /eng/acts/1/ » → 404). La contrainte « chapitre =
   LETTRE(S)-CHIFFRE » du regex exclut automatiquement ces cas (numéros nus). →
   aucun lien mort. (Extension future possible : liens d'article pinpoint
   « /section-N.html » ; non implémenté, on ne lie que le niveau loi.)

3. **Lois fédérales ANNUELLES** « S.C. 2010, c. 5 » (anglais) / « L.C. 2010,
   c. 5 » (français) → **Justice Canada** (laws-lois.justice.gc.ca). URL
   **déterministe** : le slug est « {année}_{chapitre} » (S.C. → AnnualStatutes,
   L.C. → LoisAnnuelles) :

     « S.C. 2010, c. 5 » → .../eng/AnnualStatutes/2010_5/page-1.html
     « L.C. 2010, c. 5 » → .../fra/LoisAnnuelles/2010_5/page-1.html

   Comme pour les lois révisées, l'URL est **vérifiable par requête HTTP**.
   **Limité aux années ≥ 2001** : le site ne couvre pas les lois antérieures
   (vérifié : 2000 → 404) → les citations d'année < 2001 ne sont pas liées (zéro
   lien mort). Disjoint des lois révisées (codées « 1985 », année « 20xx » ici) ;
   le lookbehind du regex évite de capter le « S.C. » de « R.S.C. ».

Point d'injection : `renderer._emit_runs`, seul endroit qui écrit les runs du
corps (prose ET blocs en retrait, donc les citations à l'intérieur d'une citation
en bloc sont couvertes), délègue à `emit_runs` ci-dessous. python-docx n'a pas
d'API hyperlien : on manipule l'OOXML (relation externe + `<w:hyperlink>`).
"""

from __future__ import annotations

import re

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Patron d'URL CanLII par cour. La langue suit la cour : SCC → /en/ca/scc,
# CSC → /fr/ca/csc (le PDF anglais cite « SCC », le français « CSC »).
_CANLII = {
    "SCC": "https://www.canlii.org/en/ca/scc/doc/{y}/{y}scc{n}/{y}scc{n}.html",
    "CSC": "https://www.canlii.org/fr/ca/csc/doc/{y}/{y}csc{n}/{y}csc{n}.html",
}
# Référence neutre CSC : groupes (année, cour, numéro). Identique à
# parser._CITATION. Le `\b...\d+\b` écarte les marqueurs « [156] » (pas de cour)
# et les années « [2017] » (pas de cour) — seul « AAAA SCC/CSC N » est capté.
_CITE_LINK = re.compile(r"\b(\d{4})\s+(SCC|CSC)\s+(\d+)\b")

# Patron d'URL Justice Canada par marqueur. R.S.C. (anglais) → /eng/acts,
# L.R.C. (français) → /fra/lois ; le chapitre cité est le slug d'URL.
_JUSTICE = {
    "RSC": "https://laws-lois.justice.gc.ca/eng/acts/{ch}/",
    "LRC": "https://laws-lois.justice.gc.ca/fra/lois/{ch}/",
}
# Loi des L.R.C. 1985 : groupes (marqueur, chapitre). Le marqueur (R.S.C. vs
# L.R.C., points facultatifs) donne la langue de l'URL ; le chapitre doit être
# LETTRE(S)-CHIFFRE (« C-50 », « F-7 », « P-1 », « C-46.1 ») — cette contrainte
# **exclut** les lois annuelles « S.C. 2012, c. 1 » et les suppléments
# « c. 1 (5e suppl.) » (chapitres nus, sans slug déterministe → non liés).
_STATUTE_LINK = re.compile(
    r"\b(R\.?S\.?C\.?|L\.?R\.?C\.?)\s+1985,?\s+c\.?\s*([A-Z]{1,2}-\d+(?:\.\d+)?)\b"
)

# Patron d'URL Justice Canada pour les lois ANNUELLES fédérales (S.C./L.C.).
# S.C. (anglais) → AnnualStatutes, L.C. (français) → LoisAnnuelles ; le slug est
# « {année}_{chapitre} ». Couverture du site à partir de 2001 (vérifié : 2000 →
# 404), d'où la borne `_ANNUAL_MIN_YEAR`.
_ANNUAL = {
    "SC": "https://laws-lois.justice.gc.ca/eng/AnnualStatutes/{y}_{n}/page-1.html",
    "LC": "https://laws-lois.justice.gc.ca/fra/LoisAnnuelles/{y}_{n}/page-1.html",
}
# Loi annuelle : groupes (marqueur, année, chapitre). Le lookbehind `(?<![.\w])`
# évite de capter le « S.C. » DANS « R.S.C. » (un point précède) ou « L.C. » dans
# « L.R.C. » (une lettre précède le L final → la séquence est L-R-C, pas L-C).
# L'année doit être « 20xx » → familles disjointes des lois révisées (codées
# « 1985 ») ; le chapitre est un entier nu (« c. 5 »), pas un slug lettre-chiffre.
_ANNUAL_LINK = re.compile(r"(?<![.\w])(S\.?C\.?|L\.?C\.?)\s+(20\d{2}),?\s+c\.?\s*(\d+)\b")
# Justice Canada ne couvre les lois annuelles qu'à partir de 2001 (avant → 404) ;
# on ne lie donc pas les années antérieures, pour éviter tout lien mort.
_ANNUAL_MIN_YEAR = 2001

# --- Style du lien (constantes uniques, faciles à changer) ------------------- #
# Défaut : bleu hyperlien Word + souligné, la convention reconnaissable
# « cliquable ». Pour un rendu sobre (noir, sans soulignement) :
#   _LINK_COLOR = None   et   _LINK_UNDERLINE = False
# La mise en forme du texte environnant (italique/gras/corps) est toujours
# préservée sur le lien (un lien dans une citation en bloc reste en 10 pt).
_LINK_COLOR = "0563C1"      # hex RVB du bleu hyperlien Word ; None = couleur du texte
_LINK_UNDERLINE = True


def canlii_url(year: str, court: str, num: str) -> str:
    """Construit l'URL CanLII d'une référence neutre (aucun appel réseau)."""
    return _CANLII[court].format(y=year, n=num)


def justice_url(marker: str, chapter: str) -> str:
    """URL Justice Canada d'une loi R.S.C./L.R.C. 1985 (aucun appel réseau).

    Le marqueur (« R.S.C. »/« L.R.C. », points facultatifs) donne la langue ;
    le chapitre cité est le slug d'URL.
    """
    key = "RSC" if marker.upper().replace(".", "").startswith("RSC") else "LRC"
    return _JUSTICE[key].format(ch=chapter)


def annual_url(marker: str, year: str, num: str) -> str | None:
    """URL Justice Canada d'une loi annuelle S.C./L.C., ou None si l'année est
    antérieure à 2001 (non couverte par le site → on ne lie pas, pour éviter un
    lien mort).

    Le marqueur (« S.C. »/« L.C. », points facultatifs) donne la langue ;
    le slug est « {année}_{chapitre} ».
    """
    if int(year) < _ANNUAL_MIN_YEAR:
        return None
    key = "SC" if marker.upper().replace(".", "").startswith("S") else "LC"
    return _ANNUAL[key].format(y=year, n=num)


def _link_spans(text):
    """(début, fin, url) de tous les liens du texte, triés, sans chevauchement.

    Collecte les trois familles (références neutres → CanLII, lois révisées 1985
    → Justice Canada, lois annuelles S.C./L.C. ≥ 2001 → Justice Canada), trie par
    position et écarte tout chevauchement (garde le premier). Un même run peut
    contenir plusieurs types.
    """
    spans = []
    for m in _CITE_LINK.finditer(text):
        year, court, num = m.groups()
        spans.append((m.start(), m.end(), canlii_url(year, court, num)))
    for m in _STATUTE_LINK.finditer(text):
        marker, chapter = m.groups()
        spans.append((m.start(), m.end(), justice_url(marker, chapter)))
    for m in _ANNUAL_LINK.finditer(text):
        marker, year, num = m.groups()
        url = annual_url(marker, year, num)
        if url:  # None si année < 2001 (non couverte) → pas de lien
            spans.append((m.start(), m.end(), url))
    spans.sort()
    out, last = [], -1
    for s, e, url in spans:
        if s >= last:  # écarte un chevauchement (garde le premier rencontré)
            out.append((s, e, url))
            last = e
    return out


def _add_hyperlink(para, url, text, *, italic=False, bold=False, size=None) -> None:
    """Ajoute un run-hyperlien à la **fin** du paragraphe.

    `para.add_run()` et `para._p.append(...)` ajoutent tous deux en fin de
    `<w:p>` : en traitant les segments de gauche à droite et en les ajoutant
    aussitôt, l'ordre d'origine est préservé. `relate_to` déduplique les
    relations externes identiques (une seule entrée .rels par URL).
    """
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if _LINK_COLOR:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), _LINK_COLOR)
        rPr.append(c)
    if _LINK_UNDERLINE:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if italic:
        rPr.append(OxmlElement("w:i"))
    if bold:
        rPr.append(OxmlElement("w:b"))
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))  # taille en demi-points
        rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")  # ne pas perdre d'espace de bord
    t.text = text
    r.append(t)
    link.append(r)
    para._p.append(link)


def emit_runs(para, runs, size=None) -> None:
    """Écrit les fragments stylés dans `para`, en transformant chaque référence
    neutre CSC en hyperlien CanLII et chaque loi R.S.C./L.R.C. 1985 en hyperlien
    Justice Canada (le reste en runs normaux), dans l'ordre.

    `size` (un `Pt`, ou None) force le corps des runs (texte en retrait, plus
    petit). Citations et lois tiennent dans un **seul** run roman (le nom de
    cause / titre de loi italique est un run distinct), donc le découpage
    intra-run suffit : on n'essaie pas de recoller une citation coupée entre deux
    runs (cas très rare).
    """
    for r in runs:
        pos = 0
        for s, e, url in _link_spans(r.text):
            before = r.text[pos:s]
            if before:
                run = para.add_run(before)
                run.italic, run.bold = r.italic, r.bold
                if size is not None:
                    run.font.size = size
            _add_hyperlink(
                para, url, r.text[s:e],
                italic=r.italic, bold=r.bold, size=size,
            )
            pos = e
        rest = r.text[pos:]
        if rest or pos == 0:  # le reste, OU le run entier s'il n'a aucun lien
            run = para.add_run(rest if pos else r.text)
            run.italic, run.bold = r.italic, r.bold
            if size is not None:
                run.font.size = size
